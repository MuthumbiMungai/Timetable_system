from docx import Document
from docx.shared import Pt
from database.db import get_connection

def export_department_timetable(department_code: str, output_file=None):
    """Export timetable in the exact format as your APPLIED SCIENCES draft"""
    if output_file is None:
        output_file = f"{department_code}_TIMETABLE_2026.docx"

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("SELECT id, name FROM departments WHERE code = ?", (department_code,))
    dept = cur.fetchone()
    if not dept:
        print("Department not found!")
        return

    doc = Document()
    doc.add_heading(f"{dept['name'].upper()} TIMETABLE 2026", level=0)

    # Get intakes for this department
    cur.execute("""
        SELECT i.id, c.name as course, i.label
        FROM intakes i
        JOIN courses c ON i.course_id = c.id
        WHERE c.department_id = (SELECT id FROM departments WHERE code=?)
        ORDER BY c.name, i.label
    """, (department_code,))
    
    for intake in cur.fetchall():
        intake_id, course, label = intake['id'], intake['course'], intake['label']

        # Course + Intake Header
        p = doc.add_paragraph()
        p.add_run(f"{course}   {label}").bold = True
        p.alignment = 1  # center

        # Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        hdr[0].text = "DAYS"
        hdr[1].text = "COURSE\nINTAKE"
        hdr[2].text = "8.00 AM – 9.30 AM"
        hdr[3].text = "10.30 AM -12.00 PM"
        hdr[4].text = "1.00 PM – 2.30 PM"
        hdr[5].text = "3.00 – 4.30 PM"

        # For now we leave the body empty (you will fill later)
        # We will improve this once you start adding timetable entries

        doc.add_paragraph()   # spacing between groups

    doc.save(output_file)
    print(f"✅ Exported: {output_file}")